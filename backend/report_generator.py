#!/usr/bin/env python3
"""
Libyana NPM - Report Generator
Builds the daily "copy-paste and send" network performance report
(text + Excel) directly from the CSV history in output/csv/.

Report structure (per spec, Phase 2 of the Unified Network Monitoring
Platform roadmap):
  1.  Executive Summary
  2.  Technology Scorecards (2G/3G/4G, busy-hour KPIs vs threshold)
  3.  Worst Cells (Top 10, combined across technologies)
  4.  Site Health & Topology Impact (basic — per-tech availability proxy;
      real alarm-to-topology correlation is Phase 3/4, once NetEco is wired
      in. The FN/HUB/Site reference table is already loaded and reported on
      so Phase 4 has something to correlate against.)
  5.  Alarms Summary — placeholder (Phase 3, pending NetEco integration)
  6.  ISP & External Traffic — placeholder (Phase 5, pending ISP feed)
  7.  Traffic & Capacity (daily traffic volumes)
  8.  Site Inventory & Availability (multi-RAT site composition)
  9.  Data Freshness (per-file latest date vs target date)
  10. Trend (last 7 days) for the headline KPI per technology

All KPI names/thresholds/weights/operators/dimensions come from
config/kpi_thresholds.csv at runtime (see HealthChecker.load_thresholds) —
nothing about a specific KPI's pass/fail limit is hardcoded here.
"""

import os
import io
import re
import logging
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use('Agg')  # headless - this runs server-side, no display available
import matplotlib.pyplot as plt

from backend.health_checker import HealthChecker

logger = logging.getLogger(__name__)

SIGNATURE_BLOCK = [
    "Faraj Ramadan Elshahaibi",
    "RF Team Leader",
    "Network Optimization Engineer / East Area",
    "Network Operation Department / East Area",
    "Libyana Mobile Phone",
    "Benghazi - Libya",
    "+218 94 7776248",
]

# Busy-hour KPI sheet used for scorecards/trend (accessibility/retainability/mobility)
SCORECARD_SHEETS = {
    'GSM': '2G_NWBH',
    'UMTS': '3G_NWBH',
    'LTE': '4G_NWBH',
}
TECH_LABELS = {'GSM': '2G (GSM)', 'UMTS': '3G (UMTS)', 'LTE': '4G (LTE)'}

# Daily traffic volume columns, per technology (not threshold-checked KPIs,
# so pulled directly from the *_NW_Daily sheets rather than kpi_thresholds.csv)
TRAFFIC_SPECS = [
    ('GSM', '2G_NW_Daily', [
        ('TCH Traffic (Erl)', 'K3014:Traffic Volume on TCH(Erl)'),
        ('SDCCH Traffic (Erl)', 'K3004:Traffic Volume on SDCCH(Erl)'),
        ('PS Traffic (MB)', 'PS Traffic (RLC)(MB)'),
    ]),
    ('UMTS', '3G_NW_Daily', [
        ('CS Traffic (Erl)', 'CS Traffic(Erl)'),
        ('PS Traffic (GB)', 'PS traffic (UL+DL)(GB)'),
    ]),
    ('LTE', '4G_NW_Daily', [
        ('DL Traffic (GB)', 'DL Traffic  Volume(GB)'),
        ('UL Traffic (GB)', 'UL Traffic  Volume(GB)'),
        ('VoLTE Traffic (Erl)', 'VoLTE Traffic Volume (Erl)'),
    ]),
]

# Additional informational KPIs (no pass/fail threshold defined in
# kpi_thresholds.csv, but useful to trend: throughput, PRB utilization,
# traffic/user volumes) tracked in the trend section alongside every
# threshold-checked busy-hour KPI for that technology. Label -> Column_Name
# in the technology's NWBH sheet.
# UMTS 'Availability' is deliberately excluded here too - see the note on
# AVAILABILITY_PROXY below (raw values are bogus large negatives).
EXTRA_TREND_KPIS = {
    'GSM': [
        ('TCH Availability (%)', 'RR307:TCH Availability(%)'),
        ('TCH Traffic (Erl)', 'K3014:Traffic Volume on TCH(Erl)'),
        ('SDCCH Traffic (Erl)', 'K3004:Traffic Volume on SDCCH(Erl)'),
        ('PS Traffic (MB)', 'PS Traffic (RLC)(MB)'),
        ('DL EGPRS Throughput (kbit/s)', 'TL9333:Average Throughput of Downlink EGPRS RLC(kbit/s)'),
    ],
    'UMTS': [
        ('PS Traffic (GB)', 'PS traffic (UL+DL)(GB)'),
        ('HSDPA Throughput/user (Kbps)', 'HSDPA Throughput per user (Local Cell)(Kbps)'),
    ],
    'LTE': [
        ('Radio Network Availability (%)', 'Radio Network Availability Rate(%)'),
        ('DL Traffic Volume (GB)', 'Downlink Traffic Volume(GB)'),
        ('UL Traffic Volume (GB)', 'Uplink Traffic Volume (GB)'),
        ('User DL Avg Throughput (Mbps)', 'User Downlink Average Throughput (Mbps)'),
        ('User UL Avg Throughput (Mbps)', 'User Uplink Average Throughput (Mbps)'),
        ('DL PRB Utilization (%)', 'DL PRB Utilizing Rate(%)'),
        ('UL PRB Utilization (%)', 'UL PRB Utilizing Rate(%)'),
        ('Attached Users (Avg)', 'L.Traffic.User.Avg'),
    ],
}

# Availability proxy per technology (no true site-level up/down feed yet,
# so busy-hour network availability is used as a stand-in).
# NOTE: 3G_NWBH's "Availability" column is excluded here — its raw values are
# large negative numbers on every single day in the source data (e.g. -12,366
# on 2026-08-16), so it is not actually a 0-100% availability figure. This
# looks like a mapping/computation bug in network_kpi_processor.py's 3G KPI
# extraction and needs investigating at the source before it can be trusted.
AVAILABILITY_PROXY = {
    'GSM': ('2G_NWBH', 'RR307:TCH Availability(%)'),
    'UMTS': None,
    'LTE': ('4G_NWBH', 'Radio Network Availability Rate(%)'),
}

FRESHNESS_FILES = {
    'SiteSummary': 'day',
    # SiteDetail's "Last Updated" tracks the last time each site's config
    # actually CHANGED, not the last pipeline run - most rows stay old on
    # purpose since bands/sectors rarely change, so it's not a freshness
    # signal and must not be treated as one here (would false-alarm daily).
    'SiteDetail': None,
    'User_Summary': 'Date',
    '2G_NWBH': 'Date', '2G_NW_Daily': 'Date', '2G_Cell_CSBH': 'Date',
    '3G_NWBH': 'Date', '3G_NW_Daily': 'Date', '3G_Cell_CSBH': 'Date',
    '4G_NWBH': 'Date', '4G_NW_Daily': 'Date', '4G_Cell_BH': 'Date',
    'Traffic_Network_2G': 'Date', 'Traffic_Network_3G': 'Date', 'Traffic_Network_4G': 'Date',
}

CELL_SHEETS = {'GSM': '2G_Cell_CSBH', 'UMTS': '3G_Cell_CSBH', 'LTE': '4G_Cell_BH'}

# The "site" grouping column is named differently per technology's cell
# sheet - 2G calls it Site Name, 3G calls it NodeB Name, 4G calls it
# eNodeB Name. Callers that want to group cells by site must resolve the
# right column per tech via this map rather than assuming 'Site Name'.
SITE_COL_BY_TECH = {'GSM': 'Site Name', 'UMTS': 'NodeB Name', 'LTE': 'eNodeB Name'}

TOPOLOGY_FILE = "config/site_topology.csv"

# Multi-RAT site composition, straight from SiteSummary.csv (section 8)
SITE_OVERLAP_COLS = [
    ('2G only sites', '2G Only'),
    ('3G only sites', '3G Only'),
    ('4G only sites', '4G Only'),
    ('2G+3G sites', '2G+3G'),
    ('2G+4G sites', '2G+4G'),
    ('3G+4G sites', '3G+4G'),
    ('2G+3G+4G sites', '2G+3G+4G (Full)'),
]

SUGGESTION_RULES = [
    ('Drop', 'Investigate interference / hardware / neighbor list'),
    ('HO', 'Review handover / neighbor parameters'),
    ('Handover', 'Review handover / neighbor parameters'),
    ('PRB', 'Evaluate cell for capacity expansion'),
    ('Congestion', 'Evaluate cell for capacity expansion'),
    ('Setup', 'Check RF resource / backhaul capacity'),
    ('RRC', 'Check RF resource / backhaul capacity'),
    ('Packet Loss', 'Check backhaul link quality'),
]


class ReportGenerator:
    """Builds the daily network performance report from output/csv/."""

    def __init__(self, csv_folder: str = "output/csv", output_folder: str = "output/reports",
                 threshold_file: str = "config/kpi_thresholds.csv"):
        self.csv_folder = csv_folder
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)
        self.health_checker = HealthChecker(threshold_file)

    # ------------------------------------------------------------------
    # Small helpers
    # ------------------------------------------------------------------

    def get_score_icon(self, score: float) -> Tuple[str, str]:
        if score >= 95:
            return "🟢", "Good"
        elif score >= 90:
            return "🟡", "Fair"
        elif score >= 80:
            return "🟠", "Poor"
        else:
            return "🔴", "Critical"

    def _load_csv(self, name: str) -> Optional[pd.DataFrame]:
        path = os.path.join(self.csv_folder, f"{name}.csv")
        if not os.path.exists(path):
            return None
        try:
            df = pd.read_csv(path)
            return df if not df.empty else None
        except Exception as e:
            logger.warning(f"Could not read {path}: {e}")
            return None

    def _row_for_date(self, df: Optional[pd.DataFrame], date_str: Optional[str],
                       date_col: str = 'Date') -> Optional[Dict]:
        if df is None or not date_str or date_col not in df.columns:
            return None
        matches = df[df[date_col].astype(str) == str(date_str)]
        if matches.empty:
            return None
        return matches.iloc[-1].to_dict()

    def _suggest_action(self, failing_kpis_text: str) -> str:
        for keyword, suggestion in SUGGESTION_RULES:
            if keyword.lower() in failing_kpis_text.lower():
                return suggestion
        return 'Investigate further'

    def get_sheet(self, name: str) -> Optional[pd.DataFrame]:
        """Public accessor for a raw CSV sheet by name (e.g. for ad-hoc dashboard queries)."""
        return self._load_csv(name)

    def get_available_dates(self, sheet: str = '4G_NWBH', date_col: str = 'Date') -> List[str]:
        """Sorted (desc) list of dates present in a Date-columned sheet, for date pickers."""
        df = self._load_csv(sheet)
        if df is None or date_col not in df.columns:
            return []
        dates = pd.to_datetime(df[date_col], errors='coerce').dropna().dt.strftime('%Y-%m-%d').unique().tolist()
        return sorted(dates, reverse=True)

    # ------------------------------------------------------------------
    # Section builders
    # ------------------------------------------------------------------

    def build_technology_scorecard(self, tech: str, target_date: str,
                                    previous_date: Optional[str] = None) -> pd.DataFrame:
        """Busy-hour KPIs vs threshold for one technology, with delta vs yesterday."""
        sheet_name = SCORECARD_SHEETS[tech]
        df = self._load_csv(sheet_name)
        today_row = self._row_for_date(df, target_date)
        prev_row = self._row_for_date(df, previous_date)

        thresholds = self.health_checker.thresholds
        if today_row is None or thresholds is None:
            return pd.DataFrame()

        tech_thresholds = thresholds[
            (thresholds['Technology'] == tech) &
            (thresholds['Source_Sheet'].str.contains(sheet_name, na=False))
        ]

        rows = []
        for _, r in tech_thresholds.iterrows():
            col = r['Column_Name']
            if col not in today_row:
                continue
            today_val = today_row.get(col)
            passed, _ = self.health_checker.check_kpi(today_val, r['Threshold'], r['Operator'])

            delta = None
            if prev_row is not None and prev_row.get(col) is not None:
                try:
                    delta = float(today_val) - float(prev_row.get(col))
                except (TypeError, ValueError):
                    delta = None

            # Gap = margin to the threshold, sign-normalized so positive
            # always means "healthy" and negative always means "failing by
            # this much", regardless of whether the operator is >= or <=.
            gap = None
            try:
                tv, th = float(today_val), float(r['Threshold'])
                gap = (tv - th) if r['Operator'] in ('>=', '>') else (th - tv)
            except (TypeError, ValueError):
                gap = None

            rows.append({
                'KPI Name': r['KPI_Name'],
                'Threshold': f"{r['Operator']} {r['Threshold']}",
                "Today's Value": today_val,
                'Status': 'PASS' if passed else 'FAIL',
                'Gap': gap,
                'Weight': r['Weight'],
                'Dimension': r['Dimension'],
                'Delta vs Yesterday': delta,
            })

        return pd.DataFrame(rows)

    def build_all_scorecards(self, target_date: str, previous_date: Optional[str]) -> Dict[str, pd.DataFrame]:
        return {tech: self.build_technology_scorecard(tech, target_date, previous_date)
                for tech in SCORECARD_SHEETS}

    def compute_health_from_scorecards(self, scorecards: Dict[str, pd.DataFrame]) -> Dict:
        """Weighted pass-rate scoring computed strictly from target_date's rows
        (replaces HealthChecker.get_health_summary, which does not filter
        whole-network KPI values by date at all)."""
        tech_scores = {}
        dimension_scores: Dict[str, Dict] = {}
        total_weight = 0.0
        weighted_sum = 0.0
        passed_count = 0
        failed_count = 0
        alerts = []

        for tech, df in scorecards.items():
            if df is None or df.empty:
                continue
            t_weighted = 0.0
            t_total = 0.0
            for _, row in df.iterrows():
                w = float(row['Weight'])
                passed = row['Status'] == 'PASS'
                dim = row['Dimension']

                total_weight += w
                weighted_sum += w * (1 if passed else 0)
                t_weighted += w * (1 if passed else 0)
                t_total += w

                if passed:
                    passed_count += 1
                else:
                    failed_count += 1
                    alerts.append({
                        'kpi': row['KPI Name'],
                        'tech': tech,
                        'value': row["Today's Value"],
                        'threshold': row['Threshold'],
                        'severity': 'Critical',
                    })

                dimension_scores.setdefault(dim, {'weighted': 0.0, 'total': 0.0})
                dimension_scores[dim]['weighted'] += w * (1 if passed else 0)
                dimension_scores[dim]['total'] += w

            tech_scores[tech] = {'score': (t_weighted / t_total * 100) if t_total else 0}

        for dim, d in dimension_scores.items():
            d['score'] = (d['weighted'] / d['total'] * 100) if d['total'] else 0

        overall_score = (weighted_sum / total_weight * 100) if total_weight else 0

        return {
            'overall_score': overall_score,
            'by_technology': tech_scores,
            'by_dimension': dimension_scores,
            'total_kpis': passed_count + failed_count,
            'passed_kpis': passed_count,
            'failed_kpis': failed_count,
            'alerts': alerts,
        }

    def build_worst_cells(self, target_date: str, n_top: int = 10) -> Dict[str, pd.DataFrame]:
        """Return the worst-performing cells per technology, one table per tech."""
        result = {}
        for tech, sheet in CELL_SHEETS.items():
            df = self._load_csv(sheet)
            if df is None:
                continue
            worst = self.health_checker.find_worst_cells(
                df, sheet, n_top=n_top, selected_date=target_date
            )
            if worst is None or worst.empty:
                continue
            worst = worst.drop(columns=['Rank', 'Violations', 'Technology'], errors='ignore')
            worst = worst.sort_values('Severity Score', ascending=False).head(n_top).reset_index(drop=True)
            worst['Suggested Action'] = worst['Failing KPIs'].apply(self._suggest_action)
            worst.insert(0, 'Rank', range(1, len(worst) + 1))
            result[tech] = worst[['Rank', 'Cell Name', 'Severity Score', 'Failing KPIs', 'Suggested Action']]

        return result

    def combine_worst_cells(self, worst_cells: Dict[str, pd.DataFrame], n_top: int = 3) -> pd.DataFrame:
        """Combine per-technology worst-cell tables into one overall top-N (used by the exec summary)."""
        combined = []
        for tech, df in (worst_cells or {}).items():
            if df is None or df.empty:
                continue
            tagged = df.drop(columns=['Rank'], errors='ignore').copy()
            tagged.insert(0, 'Technology', tech)
            combined.append(tagged)
        if not combined:
            return pd.DataFrame()
        all_worst = pd.concat(combined, ignore_index=True)
        return all_worst.sort_values('Severity Score', ascending=False).head(n_top).reset_index(drop=True)

    def build_traffic_section(self, target_date: str, previous_date: Optional[str]) -> pd.DataFrame:
        rows = []
        for tech, sheet, metrics in TRAFFIC_SPECS:
            df = self._load_csv(sheet)
            today_row = self._row_for_date(df, target_date)
            prev_row = self._row_for_date(df, previous_date)
            for label, col in metrics:
                today_val = today_row.get(col) if today_row else None
                prev_val = prev_row.get(col) if prev_row else None
                delta = None
                if today_val is not None and prev_val is not None:
                    try:
                        delta = float(today_val) - float(prev_val)
                    except (TypeError, ValueError):
                        delta = None
                rows.append({
                    'Technology': tech, 'Metric': label,
                    'Today': today_val, 'Yesterday': prev_val, 'Delta': delta,
                })
        return pd.DataFrame(rows)

    def build_site_health(self, target_date: str) -> pd.DataFrame:
        site_df = self._load_csv('SiteSummary')
        site_row = self._row_for_date(site_df, target_date, date_col='day')

        site_cols = {'GSM': '2G physical sites', 'UMTS': '3G physical sites', 'LTE': '4G physical sites'}
        rows = []
        for tech, label in TECH_LABELS.items():
            proxy = AVAILABILITY_PROXY[tech]
            avail_value = None
            if proxy is not None:
                avail_sheet, avail_col = proxy
                avail_row = self._row_for_date(self._load_csv(avail_sheet), target_date)
                avail_value = avail_row.get(avail_col) if avail_row else None
            rows.append({
                'Technology': label,
                'Physical Sites': site_row.get(site_cols[tech]) if site_row else None,
                'Availability Proxy (%)': avail_value if avail_value is not None else 'N/A (data issue, see note)',
            })
        rows.append({
            'Technology': 'TOTAL',
            'Physical Sites': site_row.get('Total Physical Sites (2G+3G+4G)') if site_row else None,
            'Availability Proxy (%)': None,
        })
        return pd.DataFrame(rows)

    def build_topology_summary(self) -> Dict:
        """Section 4 support: readiness stats for the FN/HUB→Site reference
        table. No alarm feed exists yet, so this can only report that the
        topology is loaded and ready — real impact analysis is Phase 4."""
        if not os.path.exists(TOPOLOGY_FILE):
            return {'loaded': False}
        try:
            df = pd.read_csv(TOPOLOGY_FILE)
        except Exception as e:
            logger.warning(f"Could not read {TOPOLOGY_FILE}: {e}")
            return {'loaded': False}
        if df.empty:
            return {'loaded': False}
        return {
            'loaded': True,
            'nodes': df['Node_Name'].nunique(),
            'fn_count': df[df['Node_Type'] == 'FN']['Node_Name'].nunique(),
            'hub_count': df[df['Node_Type'] == 'HUB']['Node_Name'].nunique(),
            'site_relationships': len(df),
            'regions': sorted(df['Region'].dropna().unique().tolist()),
        }

    def build_site_inventory(self, target_date: str) -> pd.DataFrame:
        """Section 8: multi-RAT site composition (from SiteSummary.csv)."""
        site_df = self._load_csv('SiteSummary')
        site_row = self._row_for_date(site_df, target_date, date_col='day')
        if site_row is None:
            return pd.DataFrame()

        rows = [{'Category': 'TOTAL', 'Site Count': site_row.get('Total Physical Sites (2G+3G+4G)')}]
        for col, label in SITE_OVERLAP_COLS:
            rows.append({'Category': label, 'Site Count': site_row.get(col)})
        return pd.DataFrame(rows)

    @staticmethod
    def _num(row, col, default=0.0):
        if row is None:
            return default
        v = row.get(col)
        if v is None or (isinstance(v, float) and pd.isna(v)):
            return default
        return v

    def build_site_summary_cards(self, target_date: str) -> Dict[str, Optional[float]]:
        """Headline site/user/traffic numbers for the dashboard's landing tab
        (SiteSummary + User_Summary + Traffic_Network_* - raw totals, not
        threshold-checked KPIs, so no config lookup applies here)."""
        g = self._num
        site_row = self._row_for_date(self._load_csv('SiteSummary'), target_date, date_col='day')
        user_row = self._row_for_date(self._load_csv('User_Summary'), target_date)
        t2g = self._row_for_date(self._load_csv('Traffic_Network_2G'), target_date)
        t3g = self._row_for_date(self._load_csv('Traffic_Network_3G'), target_date)
        t4g = self._row_for_date(self._load_csv('Traffic_Network_4G'), target_date)

        ps_users = g(user_row, '2G PS user') + g(user_row, '3G PS user') + g(user_row, '4G PS user')
        cs_users = g(user_row, '2G CS user') + g(user_row, '3G CS user')
        ps_traffic = g(t2g, '2G PS Traffic (GB)') + g(t3g, '3G PS Traffic (GB)') + g(t4g, '4G DL Traffic (GB)')
        cs_traffic = g(t2g, '2G CS Traffic (Erl)') + g(t3g, '3G CS Traffic (Erl)')

        return {
            'Total Sites (2G+3G+4G)': g(site_row, 'Total Physical Sites (2G+3G+4G)', None),
            '2G Sites': g(site_row, '2G physical sites', None),
            '3G Sites': g(site_row, '3G physical sites', None),
            '4G Sites': g(site_row, '4G physical sites', None),
            'Total PS Users': ps_users,
            'Total CS Users': cs_users,
            'VoLTE Users': g(user_row, 'VoLTE user', None),
            'Total Subscribers': g(user_row, 'Total Subscribers', None),
            'Total PS Traffic (GB)': round(ps_traffic, 2),
            'Total CS Traffic (Erl)': round(cs_traffic, 2),
        }

    def build_network_summary_block(self, target_date: str) -> List[Tuple[str, str]]:
        """The ops team's familiar 'FYI network summary' bullet list (headline
        subscriber/traffic/KPI figures), pre-formatted as (label, value)
        pairs for the report's opening section. Raw totals/busy-hour reads,
        not threshold-checked KPIs - config lookup doesn't apply here."""
        g = self._num
        user_row = self._row_for_date(self._load_csv('User_Summary'), target_date)
        t2g = self._row_for_date(self._load_csv('Traffic_Network_2G'), target_date)
        t3g = self._row_for_date(self._load_csv('Traffic_Network_3G'), target_date)
        t4g = self._row_for_date(self._load_csv('Traffic_Network_4G'), target_date)
        lte_bh = self._row_for_date(self._load_csv('4G_NWBH'), target_date)
        lte_daily = self._row_for_date(self._load_csv('4G_NW_Daily'), target_date)

        ps_users = g(user_row, '2G PS user') + g(user_row, '3G PS user') + g(user_row, '4G PS user')
        cs_subs = g(user_row, '2G CS user') + g(user_row, '3G CS user')
        ps_traffic = g(t2g, '2G PS Traffic (GB)') + g(t3g, '3G PS Traffic (GB)') + g(t4g, '4G DL Traffic (GB)')
        cs_traffic = g(t2g, '2G CS Traffic (Erl)') + g(t3g, '3G CS Traffic (Erl)')
        lte_max_users = g(lte_daily, 'L.Traffic.User.Max') or g(lte_bh, 'L.Traffic.User.Max')
        lte_traffic = g(t4g, '4G DL Traffic (GB)')

        return [
            ('Total PS users (2G+3G+4G)', f"{ps_users:,.0f}"),
            ('Total PS traffic', f"{ps_traffic:,.2f} GB"),
            ('CS subscribers (2G+3G)', f"{cs_subs:,.0f}"),
            ('Total CS traffic (2G+3G)', f"{cs_traffic:,.2f} Erlangs"),
            ('LTE maximum attached users', f"{lte_max_users:,.0f}"),
            ('Total LTE traffic', f"{lte_traffic:,.2f} GB"),
            ('Maximum number of VoLTE users', f"{g(user_row, 'VoLTE user'):,.0f}"),
            ('Number of CS roaming users', f"{g(user_row, 'Roaming CS (Almadar)'):,.0f}"),
            ('Average ping packet loss rate', f"{g(lte_bh, 'Downlink Packet Loss'):.4f}"),
            ('RRC Setup Success Rate (%)', f"{g(lte_bh, 'RRC Setup Success Rate(%)'):.4f}"),
            ('E-RAB Setup Success Rate (%)', f"{g(lte_bh, 'E-RAB Setup Success Rate'):.4f}"),
            ('Average network availability (%)', f"{g(lte_bh, 'Radio Network Availability Rate(%)'):.2f}"),
            ('Average LTE user DL throughput', f"{g(lte_bh, 'User Downlink Average Throughput (Mbps)'):.3f} Mbps"),
            ('VoLTE setup success rate (%)', f"{g(lte_bh, 'VoLTE Setup Success Rate-ZM(%)'):.3f}"),
            ('Service drop rate (%)', f"{g(lte_bh, 'Service Drop Rate (All)'):.3f}"),
        ]

    def build_data_freshness(self, target_date: str) -> pd.DataFrame:
        target_dt = pd.to_datetime(target_date)
        rows = []
        for name, date_col in FRESHNESS_FILES.items():
            df = self._load_csv(name)
            if name == 'SiteDetail':
                # Not a daily feed - "Last Updated" tracks the last time a site's
                # config actually changed, so show that date for visibility but
                # never flag lag as staleness (see comment on FRESHNESS_FILES).
                if df is not None and 'Last Updated' in df.columns:
                    latest = pd.to_datetime(df['Last Updated'], errors='coerce').max()
                    latest_str = latest.strftime('%Y-%m-%d') if pd.notna(latest) else 'N/A'
                else:
                    latest_str = 'N/A'
                status = 'ℹ️ Snapshot (config last changed)' if df is not None else '🔴 Missing'
                rows.append({'File': name, 'Latest Date': latest_str, 'Status': status})
                continue
            if not date_col:
                status = 'ℹ️ Snapshot (no date field)' if df is not None else '🔴 Missing'
                rows.append({'File': name, 'Latest Date': 'N/A', 'Status': status})
                continue
            if df is None or date_col not in df.columns:
                rows.append({'File': name, 'Latest Date': 'N/A', 'Status': '🔴 Missing'})
                continue
            latest = pd.to_datetime(df[date_col], errors='coerce').max()
            if pd.isna(latest):
                rows.append({'File': name, 'Latest Date': 'N/A', 'Status': '🔴 Missing'})
                continue
            lag = (target_dt - latest).days
            if lag <= 0:
                status = '🟢 Current'
            elif lag == 1:
                status = '🟡 1 day behind'
            else:
                status = f'🔴 {lag} days behind'
            rows.append({'File': name, 'Latest Date': latest.strftime('%Y-%m-%d'), 'Status': status})
        return pd.DataFrame(rows)

    def build_trend(self, target_date: str, days: int = 14) -> Dict[str, pd.DataFrame]:
        """Last N days of busy-hour KPI values per technology: every
        threshold-checked KPI for that tech (config-driven, from
        kpi_thresholds.csv) plus a curated set of informational KPIs
        (throughput/PRB/traffic - see EXTRA_TREND_KPIS)."""
        thresholds = self.health_checker.thresholds
        result = {}
        target_dt = pd.to_datetime(target_date)

        for tech, sheet in SCORECARD_SHEETS.items():
            df = self._load_csv(sheet)
            if df is None or 'Date' not in df.columns:
                continue

            cols_present = {}
            if thresholds is not None:
                tech_thresholds = thresholds[
                    (thresholds['Technology'] == tech) &
                    (thresholds['Source_Sheet'].str.contains(sheet, na=False))
                ]
                for _, r in tech_thresholds.iterrows():
                    if r['Column_Name'] in df.columns:
                        cols_present[r['KPI_Name']] = r['Column_Name']

            used_cols = set(cols_present.values())
            for label, col in EXTRA_TREND_KPIS.get(tech, []):
                if col in df.columns and label not in cols_present and col not in used_cols:
                    cols_present[label] = col
                    used_cols.add(col)

            if not cols_present:
                continue

            df = df.copy()
            df['_dt'] = pd.to_datetime(df['Date'], errors='coerce')
            window = df[df['_dt'] <= target_dt].sort_values('_dt').tail(days)
            if window.empty:
                continue

            trend_df = window[['Date'] + list(cols_present.values())].rename(
                columns={v: k for k, v in cols_present.items()}
            )
            result[tech] = trend_df

        return result

    def get_trend_kpi_thresholds(self, tech: str, sheet: Optional[str] = None) -> Dict[str, Tuple[float, str]]:
        """KPI label -> (threshold, operator) for the KPIs build_trend() labels
        with a pass/fail rule, so callers (e.g. the dashboard) can draw a
        threshold reference line. KPIs without a rule (EXTRA_TREND_KPIS) are
        simply absent from the returned dict.

        `sheet` defaults to the technology's busy-hour network sheet
        (SCORECARD_SHEETS); pass a cell-level sheet (CELL_SHEETS) to get the
        same KPIs/thresholds for per-cell use - kpi_thresholds.csv's
        Source_Sheet already lists both grains for each KPI."""
        thresholds = self.health_checker.thresholds
        if thresholds is None:
            return {}
        sheet = sheet or SCORECARD_SHEETS.get(tech)
        if not sheet:
            return {}
        tech_thresholds = thresholds[
            (thresholds['Technology'] == tech) &
            (thresholds['Source_Sheet'].str.contains(sheet, na=False))
        ]
        return {r['KPI_Name']: (float(r['Threshold']), r['Operator']) for _, r in tech_thresholds.iterrows()}

    def resolve_group_to_cells(self, tech: str, group_names: List[str],
                                group_col: Optional[str] = None) -> List[str]:
        """All Cell Name values belonging to the given site(s) - lets callers
        select a site (or a group of sites for a change/operation request)
        and pull every cell/sector under it for combined analysis.
        `group_col` defaults to the right site-name column for `tech`
        (SITE_COL_BY_TECH) since it differs per technology's cell sheet."""
        sheet = CELL_SHEETS.get(tech)
        df = self._load_csv(sheet)
        if df is None or not group_names:
            return []
        group_col = group_col or SITE_COL_BY_TECH.get(tech, 'Site Name')
        if group_col not in df.columns:
            return list(group_names)
        cell_col = 'Cell Name' if 'Cell Name' in df.columns else group_col
        matches = df[df[group_col].isin(group_names)][cell_col].dropna().astype(str).unique().tolist()
        return sorted(matches)

    def build_cell_trend(self, tech: str, cell_names: List[str], target_date: Optional[str] = None,
                          days: int = 14, start_date: Optional[str] = None,
                          end_date: Optional[str] = None) -> Optional[pd.DataFrame]:
        """Busy-hour KPI values for specific cells, long-format: Date, Cell,
        <KPI columns...> - one row per cell per date, so callers can plot one
        line per cell per KPI.

        Either pass `target_date` (+ optional `days`) for "last N days ending
        at target_date" (the Cell Explorer's usage), or pass an explicit
        `start_date`/`end_date` for an arbitrary custom range (Special
        Reports' usage, e.g. verifying a change request over its own window)."""
        sheet = CELL_SHEETS.get(tech)
        df = self._load_csv(sheet)
        if df is None or not cell_names:
            return None

        cell_col = 'Cell Name' if 'Cell Name' in df.columns else ('Site Name' if 'Site Name' in df.columns else None)
        if cell_col is None:
            return None

        thresholds = self.health_checker.thresholds
        col_map = {}
        if thresholds is not None:
            tech_thresholds = thresholds[
                (thresholds['Technology'] == tech) & (thresholds['Source_Sheet'].str.contains(sheet, na=False))
            ]
            for _, r in tech_thresholds.iterrows():
                if r['Column_Name'] in df.columns:
                    col_map[r['KPI_Name']] = r['Column_Name']
        if not col_map:
            return None

        sub = df[df[cell_col].isin(cell_names)].copy()
        if sub.empty:
            return None
        sub['_dt'] = pd.to_datetime(sub['Date'], errors='coerce')

        if start_date or end_date:
            if start_date:
                sub = sub[sub['_dt'] >= pd.to_datetime(start_date)]
            if end_date:
                sub = sub[sub['_dt'] <= pd.to_datetime(end_date)]
        else:
            target_dt = pd.to_datetime(target_date)
            sub = sub[sub['_dt'] <= target_dt]
            last_dates = sorted(sub['_dt'].dropna().unique())[-days:]
            sub = sub[sub['_dt'].isin(last_dates)]

        if sub.empty:
            return None

        keep = ['Date', cell_col] + list(col_map.values())
        result = sub[keep].rename(columns={v: k for k, v in col_map.items()}).rename(columns={cell_col: 'Cell'})
        return result.sort_values('Date')

    def get_cell_failing_kpis(self, tech: str, cell_names: List[str], target_date: str) -> pd.DataFrame:
        """Which threshold KPIs are failing today for the given cells, with a
        suggested fix per KPI (reuses the same SUGGESTION_RULES as the daily
        Worst Cells section)."""
        sheet = CELL_SHEETS.get(tech)
        df = self._load_csv(sheet)
        thresholds = self.health_checker.thresholds
        if df is None or thresholds is None or not cell_names:
            return pd.DataFrame()

        cell_col = 'Cell Name' if 'Cell Name' in df.columns else ('Site Name' if 'Site Name' in df.columns else None)
        if cell_col is None:
            return pd.DataFrame()

        target_norm = self.health_checker.normalize_date(target_date)
        df = df.copy()
        df['_norm_date'] = df['Date'].apply(self.health_checker.normalize_date)
        today = df[(df['_norm_date'] == target_norm) & (df[cell_col].isin(cell_names))]
        if today.empty:
            return pd.DataFrame()

        tech_thresholds = thresholds[
            (thresholds['Technology'] == tech) & (thresholds['Source_Sheet'].str.contains(sheet, na=False))
        ]

        rows = []
        for _, cell_row in today.iterrows():
            for _, r in tech_thresholds.iterrows():
                col = r['Column_Name']
                if col not in cell_row or pd.isna(cell_row[col]):
                    continue
                val = cell_row[col]
                passed, _ = self.health_checker.check_kpi(val, r['Threshold'], r['Operator'])
                if not passed:
                    rows.append({
                        'Cell': cell_row[cell_col],
                        'Failing KPI': r['KPI_Name'],
                        'Value': round(float(val), 3),
                        'Threshold': f"{r['Operator']} {r['Threshold']}",
                        'Suggested Fix': self._suggest_action(r['KPI_Name']),
                    })
        return pd.DataFrame(rows)

    def build_executive_summary(self, health: Dict, worst_cells: Dict[str, pd.DataFrame],
                                 freshness: pd.DataFrame, target_date: str) -> List[str]:
        lines = []
        score = health.get('overall_score', 0)
        icon, status = self.get_score_icon(score)
        lines.append(f"Overall Network Health ({target_date}): {score:.1f}% ({icon} {status})")

        critical_alerts = health.get('alerts', [])
        lines.append(f"Failing KPIs (busy hour, whole network): {len(critical_alerts)} of {health.get('total_kpis', 0)}")

        top3 = self.combine_worst_cells(worst_cells, n_top=3)
        if not top3.empty:
            names = "; ".join(f"{r['Cell Name']} ({r['Technology']}, score {r['Severity Score']:.0f})"
                               for _, r in top3.iterrows())
            lines.append(f"Top 3 Worst Cells: {names}")
        else:
            lines.append("Top 3 Worst Cells: none flagged")

        if freshness is not None and not freshness.empty:
            stale = freshness[freshness['Status'].str.contains('🔴|🟡')]
            if stale.empty:
                lines.append("Data Freshness: 🟢 all sources current")
            else:
                lines.append(f"Data Freshness: ⚠️ {len(stale)} source(s) behind — see Data Freshness section")

        return lines

    # ------------------------------------------------------------------
    # Rendering
    # ------------------------------------------------------------------

    def _render_table(self, df: pd.DataFrame, float_cols: Optional[List[str]] = None) -> str:
        """Render a DataFrame as a simple fixed-width text table for copy/paste."""
        if df is None or df.empty:
            return "  (no data)"
        display_df = df.copy()
        float_cols = float_cols or []
        for col in display_df.columns:
            if col in float_cols or pd.api.types.is_float_dtype(display_df[col]):
                display_df[col] = display_df[col].apply(
                    lambda v: f"{v:,.2f}" if pd.notna(v) and isinstance(v, (int, float)) else ("" if pd.isna(v) else str(v))
                )
            else:
                display_df[col] = display_df[col].apply(lambda v: "" if pd.isna(v) else str(v))

        widths = {c: max(len(c), display_df[c].str.len().max()) for c in display_df.columns}
        header = "  ".join(c.ljust(widths[c]) for c in display_df.columns)
        sep = "  ".join('-' * widths[c] for c in display_df.columns)
        body_lines = [
            "  ".join(str(row[c]).ljust(widths[c]) for c in display_df.columns)
            for _, row in display_df.iterrows()
        ]
        return "\n".join([header, sep] + body_lines)

    def generate_email_text(self, target_date: str, previous_date: str,
                             health: Dict, scorecards: Dict[str, pd.DataFrame],
                             worst_cells: Dict[str, pd.DataFrame], site_health: pd.DataFrame,
                             topology: Dict, traffic: pd.DataFrame,
                             site_inventory: pd.DataFrame, freshness: pd.DataFrame,
                             trend: Dict[str, pd.DataFrame]) -> str:
        lines = []
        lines.append("=" * 78)
        lines.append(f"📊 LIBYANA NETWORK PERFORMANCE REPORT - {target_date}")
        lines.append("=" * 78)
        lines.append("")
        lines.append("Dear Team,")
        lines.append("")
        lines.append("Please find below the daily Network Performance Summary for "
                      f"{target_date} (2G/3G/4G, EAST Region).")
        lines.append("")

        lines.append("📈 NETWORK SUMMARY")
        lines.append("-" * 78)
        for label, value in self.build_network_summary_block(target_date):
            lines.append(f"  • {label:<36}: {value}")
        lines.append("")

        lines.append("1. EXECUTIVE SUMMARY")
        lines.append("-" * 78)
        for line in self.build_executive_summary(health, worst_cells, freshness, target_date):
            lines.append(f"  • {line}")
        lines.append("")

        lines.append("2. TECHNOLOGY SCORECARDS (Busy Hour)")
        lines.append("-" * 78)
        for tech, label in TECH_LABELS.items():
            df = scorecards.get(tech)
            tech_score = health.get('by_technology', {}).get(tech, {}).get('score', 0)
            icon, status = self.get_score_icon(tech_score)
            lines.append(f"[{label}] Score: {tech_score:.1f}% ({icon} {status})")
            lines.append(self._render_table(df))
            lines.append("")

        lines.append("3. WORST CELLS (Top 10 per technology)")
        lines.append("-" * 78)
        for tech, label in TECH_LABELS.items():
            df = worst_cells.get(tech) if worst_cells else None
            lines.append(f"[{label}]")
            lines.append(self._render_table(df))
            lines.append("")

        lines.append("4. SITE HEALTH & TOPOLOGY IMPACT (basic)")
        lines.append("-" * 78)
        lines.append(self._render_table(site_health))
        lines.append("")
        if topology.get('loaded'):
            lines.append(f"  Topology reference loaded: {topology['nodes']} nodes "
                          f"({topology['fn_count']} FN, {topology['hub_count']} HUB), "
                          f"{topology['site_relationships']} site relationships mapped, "
                          f"regions: {', '.join(topology['regions'])}.")
        else:
            lines.append(f"  Topology reference not found ({TOPOLOGY_FILE}).")
        lines.append("  Alarm-to-topology impact correlation (\"FN X down -> N sites affected\") "
                      "is Phase 3/4, pending NetEco integration.")
        lines.append("")

        lines.append("5. ALARMS SUMMARY")
        lines.append("-" * 78)
        lines.append("  ⏳ Not yet available — pending Huawei NetEco alarm feed integration (Phase 3).")
        lines.append("")

        lines.append("6. ISP & EXTERNAL TRAFFIC")
        lines.append("-" * 78)
        lines.append("  ⏳ Not yet available — pending ISP peering/backbone traffic feed integration (Phase 5).")
        lines.append("")

        lines.append("7. TRAFFIC & CAPACITY")
        lines.append("-" * 78)
        lines.append(self._render_table(traffic))
        lines.append("")

        lines.append("8. SITE INVENTORY & AVAILABILITY")
        lines.append("-" * 78)
        lines.append(self._render_table(site_inventory))
        lines.append("  Site outages: pending NetEco alarm integration (Phase 3).")
        lines.append("")

        lines.append("9. DATA FRESHNESS")
        lines.append("-" * 78)
        lines.append(self._render_table(freshness))
        lines.append("")

        lines.append("10. TREND (Last 14 Days, Busy Hour)")
        lines.append("-" * 78)
        for tech, label in TECH_LABELS.items():
            tdf = trend.get(tech)
            if tdf is None or tdf.empty:
                continue
            lines.append(f"[{label}]")
            lines.append(self._render_table(tdf))
            lines.append("")

        lines.append("=" * 78)
        lines.append(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("Comparison baseline: previous day = " + previous_date)
        lines.append("=" * 78)
        lines.append("")
        lines.append("Best Regards,")
        lines.append("")
        lines.extend(SIGNATURE_BLOCK)

        return "\n".join(lines)

    def generate_excel_report(self, target_date: str, health: Dict,
                               scorecards: Dict[str, pd.DataFrame], worst_cells: Dict[str, pd.DataFrame],
                               site_health: pd.DataFrame, topology: Dict, traffic: pd.DataFrame,
                               site_inventory: pd.DataFrame, freshness: pd.DataFrame,
                               trend: Dict[str, pd.DataFrame]) -> str:
        filename = f"Network_Report_{target_date}.xlsx"
        filepath = os.path.join(self.output_folder, filename)

        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            summary_rows = [
                ['Overall Health Score (%)', round(health.get('overall_score', 0), 2)],
                ['KPIs Passed', health.get('passed_kpis', 0)],
                ['KPIs Failed', health.get('failed_kpis', 0)],
            ]
            for tech, label in TECH_LABELS.items():
                summary_rows.append([f'{label} Score (%)',
                                      round(health.get('by_technology', {}).get(tech, {}).get('score', 0), 2)])
            pd.DataFrame(summary_rows, columns=['Metric', 'Value']).to_excel(
                writer, sheet_name='Executive_Summary', index=False)

            for tech, sheet_label in [('GSM', '2G_Scorecard'), ('UMTS', '3G_Scorecard'), ('LTE', '4G_Scorecard')]:
                df = scorecards.get(tech)
                if df is not None and not df.empty:
                    df.to_excel(writer, sheet_name=sheet_label, index=False)

            for tech, sheet_label in [('GSM', 'Worst_Cells_2G'), ('UMTS', 'Worst_Cells_3G'), ('LTE', 'Worst_Cells_4G')]:
                df = worst_cells.get(tech) if worst_cells else None
                if df is not None and not df.empty:
                    df.to_excel(writer, sheet_name=sheet_label, index=False)

            if site_health is not None and not site_health.empty:
                site_health.to_excel(writer, sheet_name='Site_Health', index=False)

            topology_note = pd.DataFrame([{
                'Status': 'Reference loaded, not yet correlated to alarms (Phase 3/4)' if topology.get('loaded')
                          else 'Topology file not found',
                'Nodes': topology.get('nodes'),
                'Fiber Nodes (FN)': topology.get('fn_count'),
                'Hubs': topology.get('hub_count'),
                'Site Relationships': topology.get('site_relationships'),
                'Regions': ', '.join(topology.get('regions', [])),
            }])
            topology_note.to_excel(writer, sheet_name='Topology_Reference', index=False)
            if os.path.exists(TOPOLOGY_FILE):
                pd.read_csv(TOPOLOGY_FILE).to_excel(writer, sheet_name='Topology_Detail', index=False)

            pd.DataFrame([{'Status': 'Not yet available — pending Huawei NetEco alarm feed integration (Phase 3)'}]
                         ).to_excel(writer, sheet_name='Alarms_Summary', index=False)
            pd.DataFrame([{'Status': 'Not yet available — pending ISP peering/backbone traffic feed integration (Phase 5)'}]
                         ).to_excel(writer, sheet_name='ISP_Traffic', index=False)

            if traffic is not None and not traffic.empty:
                traffic.to_excel(writer, sheet_name='Traffic_Capacity', index=False)
            if site_inventory is not None and not site_inventory.empty:
                site_inventory.to_excel(writer, sheet_name='Site_Inventory', index=False)
            if freshness is not None and not freshness.empty:
                freshness.to_excel(writer, sheet_name='Data_Freshness', index=False)
            for tech, label in [('GSM', 'Trend_2G'), ('UMTS', 'Trend_3G'), ('LTE', 'Trend_4G')]:
                tdf = trend.get(tech)
                if tdf is not None and not tdf.empty:
                    tdf.to_excel(writer, sheet_name=label, index=False)

        logger.info(f"✅ Excel report saved: {filepath}")
        return filepath

    # ------------------------------------------------------------------
    # Word report (.docx) — formatted, ready to select-all/copy into email
    # ------------------------------------------------------------------

    @staticmethod
    def _docx_shade_cell(cell, hex_color: str):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    def _docx_add_table(self, doc: Document, df: pd.DataFrame, status_col: Optional[str] = None):
        if df is None or df.empty:
            p = doc.add_paragraph("(no data)")
            p.runs[0].italic = True
            return

        display_df = df.copy()
        for col in display_df.columns:
            if pd.api.types.is_float_dtype(display_df[col]):
                display_df[col] = display_df[col].apply(
                    lambda v: f"{v:,.2f}" if pd.notna(v) else "")
            else:
                display_df[col] = display_df[col].apply(lambda v: "" if pd.isna(v) else str(v))

        table = doc.add_table(rows=1, cols=len(display_df.columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for i, col in enumerate(display_df.columns):
            hdr_cells[i].text = str(col)
            self._docx_shade_cell(hdr_cells[i], '1F4E78')
            for p in hdr_cells[i].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)

        for _, row in display_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(display_df.columns):
                val = str(row[col])
                cells[i].text = val
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                if status_col and col == status_col:
                    if 'FAIL' in val or '🔴' in val:
                        self._docx_shade_cell(cells[i], 'F8CBCC')
                    elif 'PASS' in val or '🟢' in val:
                        self._docx_shade_cell(cells[i], 'C6EFCE')
                    elif '🟡' in val:
                        self._docx_shade_cell(cells[i], 'FFEB9C')
        doc.add_paragraph()

    def _docx_section_heading(self, doc: Document, text: str):
        h = doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    def _render_trend_chart_png(self, tdf: pd.DataFrame, kpi: str,
                                 threshold_info: Optional[Tuple[float, str]] = None) -> bytes:
        """Static PNG of one KPI's trend line, with a dashed threshold
        reference line and pass/fail line color - same logic as the
        dashboard's interactive chart, rendered for Word embedding."""
        fig, ax = plt.subplots(figsize=(4.6, 2.6), dpi=110)
        y = pd.to_numeric(tdf[kpi], errors='coerce')
        line_color = '#1f4e78'

        if threshold_info is not None:
            threshold, operator = threshold_info
            valid = y.dropna()
            if not valid.empty:
                ok, _ = self.health_checker.check_kpi(valid.iloc[-1], threshold, operator)
                line_color = '#2ca02c' if ok else '#d62728'
            ax.axhline(threshold, color='gray', linestyle='--', linewidth=1,
                       label=f'Threshold {operator} {threshold}')
            ax.legend(fontsize=6, loc='best', frameon=False)

        ax.plot(tdf['Date'], y, color=line_color, linewidth=1.8, marker='o', markersize=3)
        ax.set_title(kpi, fontsize=9)
        ax.tick_params(axis='x', labelrotation=45, labelsize=6)
        ax.tick_params(axis='y', labelsize=7)
        ax.grid(True, alpha=0.3)
        fig.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format='png')
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    def _render_multi_cell_trend_chart_png(self, cell_trend_df: pd.DataFrame, kpi: str,
                                            threshold_info: Optional[Tuple[float, str]] = None) -> bytes:
        """Same style as _render_trend_chart_png, but one line per 'Cell' in
        cell_trend_df - for Special Reports covering a hand-picked group of
        sites/cells, so each one is visible on the same chart."""
        fig, ax = plt.subplots(figsize=(4.6, 2.6), dpi=110)

        if threshold_info is not None:
            threshold, operator = threshold_info
            ax.axhline(threshold, color='gray', linestyle='--', linewidth=1,
                       label=f'Threshold {operator} {threshold}')

        for cell, cdf in cell_trend_df.groupby('Cell'):
            y = pd.to_numeric(cdf[kpi], errors='coerce')
            ax.plot(cdf['Date'], y, linewidth=1.6, marker='o', markersize=2.5, label=str(cell))

        ax.set_title(kpi, fontsize=9)
        ax.tick_params(axis='x', labelrotation=45, labelsize=6)
        ax.tick_params(axis='y', labelsize=7)
        ax.grid(True, alpha=0.3)
        ax.legend(fontsize=5.5, loc='best', frameon=False, ncol=2)
        fig.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format='png')
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    def _docx_add_chart_grid(self, doc: Document, tdf: pd.DataFrame, kpi_list: List[str],
                              kpi_thresholds: Dict[str, Tuple[float, str]], cols: int = 2,
                              renderer=None):
        """Lay out one PNG chart per KPI in kpi_list into a `cols`-wide grid,
        using a borderless docx table as the grid container. `renderer`
        defaults to the single-series chart; pass
        _render_multi_cell_trend_chart_png for a multi-cell comparison chart."""
        if not kpi_list:
            return
        renderer = renderer or self._render_trend_chart_png
        rows = (len(kpi_list) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        for i, kpi in enumerate(kpi_list):
            cell = table.rows[i // cols].cells[i % cols]
            png = renderer(tdf, kpi, kpi_thresholds.get(kpi))
            run = cell.paragraphs[0].add_run()
            run.add_picture(io.BytesIO(png), width=Inches(3.1))
        doc.add_paragraph()

    def generate_word_report(self, target_date: str, previous_date: str,
                              health: Dict, scorecards: Dict[str, pd.DataFrame],
                              worst_cells: Dict[str, pd.DataFrame], site_health: pd.DataFrame,
                              topology: Dict, traffic: pd.DataFrame,
                              site_inventory: pd.DataFrame, freshness: pd.DataFrame,
                              trend: Dict[str, pd.DataFrame]) -> str:
        doc = Document()
        for style_name in ('Normal',):
            style = doc.styles[style_name]
            style.font.name = 'Calibri'
            style.font.size = Pt(10)

        title = doc.add_heading('LIBYANA NETWORK PERFORMANCE REPORT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        sub = doc.add_paragraph(f"{target_date}  |  Region: EAST")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(12)
        sub.runs[0].font.bold = True
        doc.add_paragraph()

        doc.add_paragraph("Dear Team,")
        doc.add_paragraph(
            f"Please find below the daily Network Performance Summary for {target_date} "
            "(2G/3G/4G, EAST Region)."
        )

        # Site summary (Date | Region | On Air Sites | 2G | 3G | 4G)
        self._docx_section_heading(doc, 'Site Summary')
        cards = self.build_site_summary_cards(target_date)
        site_row_df = pd.DataFrame([{
            'Date': target_date,
            'Region': 'Libyana EAST Network',
            'On Air Sites': cards.get('Total Sites (2G+3G+4G)'),
            '2G': cards.get('2G Sites'),
            '3G': cards.get('3G Sites'),
            '4G': cards.get('4G Sites'),
        }])
        self._docx_add_table(doc, site_row_df)

        # Headline network summary (the ops team's familiar "FYI" bullet list)
        self._docx_section_heading(doc, 'Network Summary')
        for label, value in self.build_network_summary_block(target_date):
            doc.add_paragraph(f"{label}: {value}", style='List Bullet')

        # 1. Executive Summary
        self._docx_section_heading(doc, '1. Executive Summary')
        for line in self.build_executive_summary(health, worst_cells, freshness, target_date):
            doc.add_paragraph(line, style='List Bullet')

        # 2. Technology Scorecards
        self._docx_section_heading(doc, '2. Technology Scorecards (Busy Hour)')
        for tech, label in TECH_LABELS.items():
            tech_score = health.get('by_technology', {}).get(tech, {}).get('score', 0)
            icon, status = self.get_score_icon(tech_score)
            doc.add_heading(f"{label} — Score: {tech_score:.1f}% ({icon} {status})", level=2)
            self._docx_add_table(doc, scorecards.get(tech), status_col='Status')

        # 3. Worst Cells
        self._docx_section_heading(doc, '3. Worst Cells (Top 10 per technology)')
        for tech, label in TECH_LABELS.items():
            doc.add_heading(label, level=2)
            self._docx_add_table(doc, worst_cells.get(tech) if worst_cells else None)

        # 4. Site Health & Topology Impact
        self._docx_section_heading(doc, '4. Site Health & Topology Impact (basic)')
        self._docx_add_table(doc, site_health)
        if topology.get('loaded'):
            doc.add_paragraph(
                f"Topology reference loaded: {topology['nodes']} nodes "
                f"({topology['fn_count']} FN, {topology['hub_count']} HUB), "
                f"{topology['site_relationships']} site relationships mapped, "
                f"regions: {', '.join(topology['regions'])}."
            )
        else:
            doc.add_paragraph(f"Topology reference not found ({TOPOLOGY_FILE}).")
        note = doc.add_paragraph(
            'Alarm-to-topology impact correlation ("FN X down -> N sites affected") '
            'is Phase 3/4, pending NetEco integration.'
        )
        note.runs[0].italic = True

        # 5. Alarms Summary
        self._docx_section_heading(doc, '5. Alarms Summary')
        p = doc.add_paragraph('⏳ Not yet available — pending Huawei NetEco alarm feed integration (Phase 3).')
        p.runs[0].italic = True

        # 6. ISP & External Traffic
        self._docx_section_heading(doc, '6. ISP & External Traffic')
        p = doc.add_paragraph('⏳ Not yet available — pending ISP peering/backbone traffic feed integration (Phase 5).')
        p.runs[0].italic = True

        # 7. Traffic & Capacity
        self._docx_section_heading(doc, '7. Traffic & Capacity')
        self._docx_add_table(doc, traffic)

        # 8. Site Inventory & Availability
        self._docx_section_heading(doc, '8. Site Inventory & Availability')
        self._docx_add_table(doc, site_inventory)
        p = doc.add_paragraph('Site outages: pending NetEco alarm integration (Phase 3).')
        p.runs[0].italic = True

        # 9. Data Freshness
        self._docx_section_heading(doc, '9. Data Freshness')
        self._docx_add_table(doc, freshness, status_col='Status')

        # 10. Trend - charts, not a plain table, so pass/fail is visible at a glance
        self._docx_section_heading(doc, '10. Trend (Last 14 Days, Busy Hour)')
        for tech, label in TECH_LABELS.items():
            tdf = trend.get(tech)
            if tdf is None or tdf.empty:
                continue
            doc.add_heading(label, level=2)
            kpi_thresholds = self.get_trend_kpi_thresholds(tech)
            kpi_cols = [c for c in tdf.columns if c != 'Date']
            threshold_kpis = [c for c in kpi_cols if c in kpi_thresholds]
            other_kpis = [c for c in kpi_cols if c not in kpi_thresholds]

            if threshold_kpis:
                doc.add_heading('Threshold KPIs', level=3)
                self._docx_add_chart_grid(doc, tdf, threshold_kpis, kpi_thresholds)
            if other_kpis:
                doc.add_heading('Traffic & Capacity', level=3)
                self._docx_add_chart_grid(doc, tdf, other_kpis, kpi_thresholds)

        doc.add_paragraph('_' * 60)
        footer = doc.add_paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
                                    f"Comparison baseline: previous day = {previous_date}")
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_paragraph()
        doc.add_paragraph('Best Regards,')
        doc.add_paragraph()
        for line in SIGNATURE_BLOCK:
            doc.add_paragraph(line)

        filename = f"Network_Report_{target_date}.docx"
        filepath = os.path.join(self.output_folder, filename)
        doc.save(filepath)
        logger.info(f"✅ Word report saved: {filepath}")
        return filepath

    def generate_tables_word_report(self, title: str, tables: List[Tuple[str, Optional[pd.DataFrame]]],
                                     subtitle: str = "") -> bytes:
        """Lightweight standalone Word export for a single dashboard tab/view -
        a title plus one or more named tables, no charts. For per-tab exports
        of table-only views (Worst Cells, Scorecards, Site Inventory, etc.)
        where the full daily report or a Special Report would be overkill."""
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        if subtitle:
            sub = doc.add_paragraph(subtitle)
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.runs[0].font.bold = True
        gen = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen.runs[0].font.size = Pt(8)
        gen.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        doc.add_paragraph()

        for name, df in tables:
            self._docx_section_heading(doc, name)
            status_col = 'Status' if df is not None and 'Status' in df.columns else None
            self._docx_add_table(doc, df, status_col=status_col)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ------------------------------------------------------------------
    # Special Reports - ad-hoc report for a hand-picked group of sites/cells
    # over a custom date range (e.g. verifying a change/operation request
    # that only touched specific sites).
    # ------------------------------------------------------------------

    def generate_group_word_report(self, tech: str, cell_names: List[str], group_label: str,
                                    start_date: str, end_date: str) -> Optional[str]:
        trend = self.build_cell_trend(tech, cell_names, start_date=start_date, end_date=end_date)
        if trend is None or trend.empty:
            return None

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        title = doc.add_heading('LIBYANA SPECIAL REPORT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        sub = doc.add_paragraph(f"{group_label}  |  {TECH_LABELS.get(tech, tech)}  |  {start_date} to {end_date}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(12)
        sub.runs[0].font.bold = True
        doc.add_paragraph()

        doc.add_paragraph(f"Cells/sites included ({len(cell_names)}): " + ', '.join(sorted(cell_names)))

        latest_date = trend['Date'].max()
        self._docx_section_heading(doc, f'Combined KPIs — {latest_date}')
        latest = trend[trend['Date'] == latest_date].drop(columns=['Date'])
        self._docx_add_table(doc, latest)

        self._docx_section_heading(doc, 'Failing KPIs & Suggested Fixes')
        failing = self.get_cell_failing_kpis(tech, cell_names, latest_date)
        if failing is not None and not failing.empty:
            self._docx_add_table(doc, failing)
        else:
            doc.add_paragraph('No threshold KPIs are failing for this group on the latest date in range.')

        self._docx_section_heading(doc, f'Trend ({start_date} to {end_date})')
        kpi_thresholds = self.get_trend_kpi_thresholds(tech, sheet=CELL_SHEETS[tech])
        kpi_cols = [c for c in trend.columns if c not in ('Date', 'Cell')]
        self._docx_add_chart_grid(doc, trend, kpi_cols, kpi_thresholds,
                                   renderer=self._render_multi_cell_trend_chart_png)

        doc.add_paragraph('_' * 60)
        footer = doc.add_paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_paragraph()
        doc.add_paragraph('Best Regards,')
        doc.add_paragraph()
        for line in SIGNATURE_BLOCK:
            doc.add_paragraph(line)

        safe_label = re.sub(r'[^A-Za-z0-9_-]+', '_', group_label).strip('_')[:40] or 'Group'
        filename = f"Special_Report_{safe_label}_{start_date}_to_{end_date}.docx"
        filepath = os.path.join(self.output_folder, filename)
        doc.save(filepath)
        logger.info(f"✅ Special report saved: {filepath}")
        return filepath

    # ------------------------------------------------------------------
    # Entry point
    # ------------------------------------------------------------------

    def generate_report(self, target_date: str, previous_date: Optional[str] = None) -> Tuple[str, str, str]:
        """Build the complete daily report (text + Excel) for target_date."""
        if previous_date is None:
            previous_date = (datetime.strptime(target_date, '%Y-%m-%d') - timedelta(days=1)).strftime('%Y-%m-%d')

        scorecards = self.build_all_scorecards(target_date, previous_date)
        health = self.compute_health_from_scorecards(scorecards)
        worst_cells = self.build_worst_cells(target_date)
        site_health = self.build_site_health(target_date)
        topology = self.build_topology_summary()
        traffic = self.build_traffic_section(target_date, previous_date)
        site_inventory = self.build_site_inventory(target_date)
        freshness = self.build_data_freshness(target_date)
        trend = self.build_trend(target_date, days=14)

        email_text = self.generate_email_text(
            target_date, previous_date, health, scorecards, worst_cells, site_health,
            topology, traffic, site_inventory, freshness, trend
        )
        text_file = os.path.join(self.output_folder, f"Network_Report_{target_date}.txt")
        with open(text_file, 'w', encoding='utf-8') as f:
            f.write(email_text)

        excel_file = self.generate_excel_report(
            target_date, health, scorecards, worst_cells, site_health, topology,
            traffic, site_inventory, freshness, trend
        )

        word_file = self.generate_word_report(
            target_date, previous_date, health, scorecards, worst_cells, site_health,
            topology, traffic, site_inventory, freshness, trend
        )

        return text_file, excel_file, word_file
